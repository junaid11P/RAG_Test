import re
import fitz  # PyMuPDF
import docx
import os
import io
from typing import List

# Optional advanced libraries
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

try:
    from markitdown import MarkItDown
except ImportError:
    MarkItDown = None

class DocumentProcessor:
    _md = MarkItDown() if MarkItDown else None
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        """Extracts text from a PDF with layout awareness and table support."""
        full_text = []

        # 1. Use pdfplumber for better table extraction if available
        if pdfplumber:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        # Extract basic text
                        page_text = page.extract_text(layout=True)
                        if page_text:
                            full_text.append(page_text)
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                full_text.append("\n--- PDF Table Start ---")
                                for row in table:
                                    # Filter None values from rows
                                    row_data = [str(cell) if cell is not None else "" for cell in row]
                                    full_text.append(" | ".join(row_data))
                                full_text.append("--- PDF Table End ---\n")
                return "\n".join(full_text)
            except Exception as e:
                print(f"pdfplumber failed, falling back to PyMuPDF: {e}")

        # 2. Fallback to PyMuPDF (layout aware blocks)
        doc = fitz.open(pdf_path)
        text_blocks = []
        for page in doc:
            # get_text("blocks") preserves some structural integrity
            blocks = page.get_text("blocks")
            for block in blocks:
                # block[4] is the text content
                text_blocks.append(block[4])
        return "\n".join(text_blocks)

    @staticmethod
    def extract_text_from_docx(docx_path: str) -> str:
        """Extracts text and table data from a DOCX file."""
        doc = docx.Document(docx_path)
        content = []
        
        # 1. Process by Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 2. Process Tables (Preserving structure)
        for table in doc.tables:
            content.append("\n--- DOCX Table Start ---")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                content.append(" | ".join(row_data))
            content.append("--- DOCX Table End ---\n")
            
        return "\n".join(content)

    @staticmethod
    def extract_text_from_image(image_path: str) -> str:
        """Performs OCR on an image file."""
        if not Image or not pytesseract:
            return "[OCR Error: PIL or pytesseract not installed. Image processing skipped.]"
        
        try:
            img = Image.open(image_path)
            # Basic OCR
            text = pytesseract.image_to_string(img)
            return text
        except Exception as e:
            return f"[OCR Error: {str(e)}]"

    @staticmethod
    def extract_text_from_txt(txt_path: str) -> str:
        """Extracts text from a TXT file."""
        with open(txt_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def basic_clean(text: str) -> str:
        """Cleans text while preserving structural characters like | for tables."""
        # Keep ASCII characters + some table markers
        text = re.sub(r"[^\x00-\x7F]+", " ", text)
        # Normalize horizontal spaces but keep structure
        text = re.sub(r"[ \t]+", " ", text)
        return text.strip()

    @staticmethod
    def remove_noise(text: str) -> str:
        """Removes common digital noise without destroying table structure."""
        noise_patterns = [
            r"mailto:.*",
            r"https?://\S+",
            r"www\.\S+",
            r"copyright\s+©\s+\d{4}",
            r"page\s+\d+\s+of\s+\d+",
            r"(?i)confidential",
            r"-----------------+"
        ]
        for pattern in noise_patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE)
        return text.strip()

    @classmethod
    def process_document(cls, file_path: str) -> str:
        """Complete production pipeline: Table & Layout extraction + optional OCR."""
        ext = os.path.splitext(file_path)[1].lower()
        
        # 1. Try MarkItDown first
        if cls._md:
            try:
                result = cls._md.convert(file_path)
                if result and result.text_content:
                    # Clean the markdown content
                    clean_text = cls.basic_clean(result.text_content)
                    return cls.remove_noise(clean_text)
            except Exception as e:
                print(f"MarkItDown failed for {file_path}: {e}. Falling back to legacy extractors.")

        # 2. Legacy Fallback
        raw_text = ""
        if ext == '.pdf':
            raw_text = cls.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            raw_text = cls.extract_text_from_docx(file_path)
        elif ext == '.txt':
            raw_text = cls.extract_text_from_txt(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            raw_text = cls.extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
        clean_text = cls.basic_clean(raw_text)
        final_text = cls.remove_noise(clean_text)
        
        if not final_text.strip():
            return "No readable text found in document. If this is a scanned PDF or Image, ensure Tesseract OCR is installed."
            
        return final_text
